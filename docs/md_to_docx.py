#!/usr/bin/env python3
"""Convert docs/forth_tutorial.md to a properly styled Word document.

pandoc isn't available in this environment, so this is a small,
purpose-built Markdown -> .docx converter covering exactly the
constructs the tutorial uses: ATX headings, fenced code blocks,
pipe tables, bullet lists, images, horizontal rules, and inline
`code` / **bold** / *italic* / [links].

Usage:  python3 docs/md_to_docx.py [in.md] [out.docx]
"""

import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CODE_FONT = "Consolas"
CODE_FALLBACK = "Courier New"
CODE_SHADE = "F2F2F2"       # light grey block background
CODE_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
INLINE_SHADE = "EFEFEF"


# --------------------------------------------------------------------------
# low-level docx helpers
# --------------------------------------------------------------------------

def shade(element, fill):
    """Apply a solid background fill to a paragraph or table cell."""
    pr = element.get_or_add_pPr() if element.tag.endswith('}p') else element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pr.append(shd)


def shade_paragraph(par, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    par._p.get_or_add_pPr().append(shd)


def shade_run(run, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    run._r.get_or_add_rPr().append(shd)


def border_paragraph(par, color="D0D0D0"):
    """Thin box around a paragraph, used as the code-block border."""
    pbdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '4')
        e.set(qn('w:color'), color)
        pbdr.append(e)
    par._p.get_or_add_pPr().append(pbdr)


def no_space_after(par):
    par.paragraph_format.space_after = Pt(0)


def code_run(par, text):
    run = par.add_run(text)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = CODE_COLOR
    # east-asian + hAnsi fallbacks so Word doesn't substitute a serif
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), CODE_FONT)
    return run


# --------------------------------------------------------------------------
# inline markdown
# --------------------------------------------------------------------------

INLINE_RE = re.compile(
    r'(`[^`]+`)'                      # inline code
    r'|(\*\*.+?\*\*)'                 # bold (may itself contain `code`)
    r'|(!?\[[^\]]*\]\([^)]*\))'       # link / image
    r'|(\*[^*\s][^*]*\*)'             # italic
)


def add_inline(par, text, base_bold=False):
    """Render one line of inline markdown into an existing paragraph."""
    text = text.replace('\\|', '|')
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            plain(par, text[pos:m.start()], base_bold)
        tok = m.group(0)
        if tok.startswith('`'):
            r = code_run(par, tok[1:-1])
            shade_run(r, INLINE_SHADE)
            r.bold = base_bold
        elif tok.startswith('**'):
            # bold content can hold inline code of its own — recurse
            add_inline(par, tok[2:-2], base_bold=True)
        elif tok.startswith('[') or tok.startswith('!['):
            label = re.match(r'!?\[([^\]]*)\]', tok).group(1)
            target = re.search(r'\(([^)]*)\)', tok).group(1)
            # the label may itself be `code` or **bold** — render it properly,
            # then style whichever runs that produced
            before = len(par.runs)
            add_inline(par, label, base_bold)
            if not target.startswith('#'):
                for r in par.runs[before:]:
                    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                    r.underline = True
        else:
            r = par.add_run(tok[1:-1])
            r.italic = True
            r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        plain(par, text[pos:], base_bold)


def plain(par, text, bold=False):
    if not text:
        return
    r = par.add_run(text)
    r.bold = bold


# --------------------------------------------------------------------------
# block parsing
# --------------------------------------------------------------------------

def split_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|') and not line.endswith('\\|'):
        line = line[:-1]
    # split on unescaped pipes
    cells, cur, i = [], '', 0
    while i < len(line):
        if line[i] == '\\' and i + 1 < len(line) and line[i + 1] == '|':
            cur += '\\|'
            i += 2
            continue
        if line[i] == '|':
            cells.append(cur.strip())
            cur = ''
            i += 1
            continue
        cur += line[i]
        i += 1
    cells.append(cur.strip())
    return cells


def is_divider(line):
    return bool(re.fullmatch(r'\|[\s:|-]+\|', line.strip()))


# --------------------------------------------------------------------------
# document furniture: title page, TOC field, running header/footer
# --------------------------------------------------------------------------

def strip_md_inline(text):
    """Strip `code`/**bold**/*italic* markup, for plain-text contexts (the
    title page, the running header) where we want the words, not markup."""
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    return text


def add_field(paragraph, instr_text, placeholder=''):
    """Insert a real OOXML field (begin / instrText / separate / result / end
    run sequence) into `paragraph`. `placeholder` is the cached display text
    shown until Word (re)computes the field -- on open if update-fields-on-open
    is set, or via F9 / right-click -> Update Field otherwise. Returns the
    result run so the caller can style it.
    """
    r_begin = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r_begin._r.append(fld_begin)

    r_instr = paragraph.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    r_instr._r.append(instr)

    r_sep = paragraph.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    r_sep._r.append(fld_sep)

    r_result = paragraph.add_run(placeholder)

    r_end = paragraph.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r_end._r.append(fld_end)

    return r_result


def build_title_page(doc, title):
    """A dedicated first page: the document title and a short, honest
    one-line description -- no invented authorship."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x36, 0x4D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(14)
    r2 = sub.add_run('A from-scratch tutorial for the 2068-Forth interpreter')
    r2.font.size = Pt(14)
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    doc.add_page_break()


def build_toc_page(doc):
    """A 'Table of Contents' page holding a real TOC field.

    The heading here is deliberately built with direct formatting, not the
    Heading1/Heading2 styles, so the `\\o "1-3"` TOC field doesn't list
    itself as an entry.
    """
    heading = doc.add_paragraph()
    r = heading.add_run('Table of Contents')
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x36, 0x4D)
    heading.paragraph_format.space_after = Pt(12)

    toc_par = doc.add_paragraph()
    add_field(
        toc_par,
        'TOC \\o "1-3" \\h \\z \\u',
        'Right-click here and choose "Update Field" (or press F9) to '
        'build the table of contents.',
    )

    doc.add_page_break()


def configure_header_footer(doc, title):
    """Running header (doc title) + footer (page number) on every page
    after the title page. Since `different_first_page_header_footer` only
    suppresses page 1's header/footer, the title page stays clean while the
    TOC page and every section page onward get both.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header_par = section.header.paragraphs[0]
    header_par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_par.paragraph_format.space_after = Pt(0)
    hr = header_par.add_run(title)
    hr.font.size = Pt(9)
    hr.italic = True
    hr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), '4')
    b.set(qn('w:space'), '4')
    b.set(qn('w:color'), 'BFBFBF')
    pbdr.append(b)
    header_par._p.get_or_add_pPr().append(pbdr)

    footer_par = section.footer.paragraphs[0]
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_par.add_run('Page ')
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    result_run = add_field(footer_par, 'PAGE', '1')
    result_run.font.size = Pt(9)
    result_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


# `w:updateFields` must land in the exact schema position python-docx's own
# CT_Settings sequence expects (right before hdrShapeDefaults/compat/etc.,
# after alwaysMergeEmptyNamespace) or Word will consider the part invalid
# and "repair" the file on open.
_UPDATE_FIELDS_SUCCESSORS = (
    'w:hdrShapeDefaults', 'w:footnotePr', 'w:endnotePr', 'w:compat',
    'w:docVars', 'w:rsids', 'm:mathPr', 'w:attachedSchema',
    'w:themeFontLang', 'w:clrSchemeMapping', 'w:doNotIncludeSubdocsInStats',
    'w:doNotAutoCompressPictures', 'w:forceUpgrade', 'w:captions',
    'w:readModeInkLockDown', 'w:smartTagType', 'sl:schemaLibrary',
    'w:shapeDefaults', 'w:doNotEmbedSmartTags', 'w:decimalSymbol',
    'w:listSeparator',
)


def enable_update_fields_on_open(doc):
    """Set settings.xml's w:updateFields so Word recalculates all fields
    (the TOC included) automatically the first time the document is opened,
    instead of requiring a manual Update Field / F9.
    """
    settings = doc.settings.element
    if settings.find(qn('w:updateFields')) is not None:
        return
    el = OxmlElement('w:updateFields')
    el.set(qn('w:val'), 'true')
    successor_qns = [qn(tag) for tag in _UPDATE_FIELDS_SUCCESSORS]
    for child in list(settings):
        if child.tag in successor_qns:
            child.addprevious(el)
            return
    settings.append(el)


def convert(md_path, docx_path):
    src_dir = os.path.dirname(os.path.abspath(md_path))
    lines = open(md_path, encoding='utf-8').read().split('\n')

    doc = Document()

    # base body font
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)

    # ---- title page + Table of Contents -----------------------------------
    # The .md's single top-level `# Title` heading becomes the title page
    # and the running header text, rather than an in-body Heading1 -- the
    # body's own headings (`## N. Section`, `### Subsection`) start counting
    # from Heading1 themselves (see the heading handler below).
    doc_title = 'Learning Forth on 2068-Forth'
    i = 0
    n = len(lines)
    if lines and re.match(r'^#\s+\S', lines[0].strip()):
        doc_title = strip_md_inline(
            re.match(r'^#\s+(.*)$', lines[0].strip()).group(1)
        )
        i = 1
    build_title_page(doc, doc_title)
    build_toc_page(doc)

    para_buf = []
    in_list = False

    def flush_paragraph():
        nonlocal para_buf
        if not para_buf:
            return
        text = ' '.join(s.strip() for s in para_buf).strip()
        para_buf = []
        if not text:
            return
        p = doc.add_paragraph()
        add_inline(p, text)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # ---- fenced code block -------------------------------------------
        if stripped.startswith('```'):
            flush_paragraph()
            i += 1
            body = []
            while i < n and not lines[i].strip().startswith('```'):
                body.append(lines[i].rstrip())
                i += 1
            i += 1  # closing fence
            while body and not body[0].strip():
                body.pop(0)
            while body and not body[-1].strip():
                body.pop()
            for idx, code_line in enumerate(body):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6 if idx == 0 else 0)
                p.paragraph_format.space_after = Pt(6 if idx == len(body) - 1 else 0)
                p.paragraph_format.line_spacing = 1.0
                shade_paragraph(p, CODE_SHADE)
                code_run(p, code_line if code_line else ' ')
            continue

        # ---- heading ------------------------------------------------------
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            flush_paragraph()
            # The .md's single `#` (level 1) document title was already
            # consumed into the title page/header above and never reaches
            # this loop, so body headings start one level up: `##` (the
            # section headings) map to real Word Heading1, `###` to
            # Heading2, and so on -- this is what lets `## N. Section Name`
            # (and `### Subsection`) show up as real, TOC-navigable
            # headings rather than just bold/large text.
            level = max(1, len(m.group(1)) - 1)
            p = doc.add_heading(level=min(level, 4))
            for r in list(p.runs):
                r.text = ''
            add_inline(p, m.group(2))
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x1F, 0x36, 0x4D)
            i += 1
            continue

        # ---- horizontal rule ---------------------------------------------
        if re.fullmatch(r'-{3,}', stripped):
            flush_paragraph()
            p = doc.add_paragraph()
            no_space_after(p)
            pbdr = OxmlElement('w:pBdr')
            b = OxmlElement('w:bottom')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '6')
            b.set(qn('w:space'), '1')
            b.set(qn('w:color'), 'BFBFBF')
            pbdr.append(b)
            p._p.get_or_add_pPr().append(pbdr)
            i += 1
            continue

        # ---- image (standalone line) -------------------------------------
        im = re.fullmatch(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if im:
            flush_paragraph()
            alt, rel = im.group(1), im.group(2)
            path = os.path.join(src_dir, rel)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(4.2))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(cap, alt)
                for r in cap.runs:
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            else:
                p = doc.add_paragraph()
                add_inline(p, alt)
            i += 1
            continue

        # ---- pipe table ---------------------------------------------------
        if stripped.startswith('|') and i + 1 < n and is_divider(lines[i + 1]):
            flush_paragraph()
            header = split_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(split_row(lines[i]))
                i += 1
            ncols = len(header)
            table = doc.add_table(rows=1, cols=ncols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = table.rows[0].cells
            for c, txt in enumerate(header):
                par = hdr[c].paragraphs[0]
                par.paragraph_format.space_after = Pt(2)
                add_inline(par, txt, base_bold=True)
                shade(hdr[c]._tc.get_or_add_tcPr(), 'DCE6F1')
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncols):
                    txt = row[c] if c < len(row) else ''
                    par = cells[c].paragraphs[0]
                    par.paragraph_format.space_after = Pt(2)
                    add_inline(par, txt)
            doc.add_paragraph()
            continue

        # A list marker only counts at the start of a block — otherwise a
        # wrapped prose line such as "2. Device 1 reports ..." would be
        # mistaken for a numbered item.
        block_start = (not para_buf) or in_list

        # ---- bullet list ---------------------------------------------------
        if block_start and re.match(r'^[-*]\s+\S', stripped):
            flush_paragraph()
            item = re.sub(r'^[-*]\s+', '', stripped)
            i += 1
            # continuation lines are indented
            while i < n and lines[i].startswith('  ') and lines[i].strip() \
                    and not re.match(r'^\s*[-*]\s+', lines[i]):
                item += ' ' + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, item)
            in_list = True
            continue

        # ---- numbered list --------------------------------------------------
        if block_start and re.match(r'^\d+\.\s+\S', stripped):
            flush_paragraph()
            item = re.sub(r'^\d+\.\s+', '', stripped)
            i += 1
            while i < n and lines[i].startswith('  ') and lines[i].strip() \
                    and not re.match(r'^\s*(\d+\.|[-*])\s+', lines[i]):
                item += ' ' + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style='List Number')
            add_inline(p, item)
            in_list = True
            continue

        # ---- blank line / ordinary prose -------------------------------------
        if not stripped:
            flush_paragraph()
            in_list = False
        else:
            para_buf.append(stripped)
            in_list = False
        i += 1

    flush_paragraph()

    configure_header_footer(doc, doc_title)
    enable_update_fields_on_open(doc)

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    here = os.path.dirname(os.path.abspath(__file__))
    src = sys.argv[1] if len(sys.argv) > 1 else os.path.join(here, 'forth_tutorial.md')
    dst = sys.argv[2] if len(sys.argv) > 2 else os.path.join(here, 'forth_tutorial.docx')
    print('wrote', convert(src, dst))
